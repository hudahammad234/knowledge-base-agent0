"""
loader.py
Handles automatic loading/indexing of documents: PDF, DOCX, TXT, Markdown, CSV.
Detects newly added documents (via content hash) so re-indexing does not
require rebuilding the whole project.

Member responsible: Member 1 - Knowledge Base
"""

import os
import csv
import hashlib
from dataclasses import dataclass, field
from typing import List, Dict

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    import docx  # python-docx
except ImportError:
    docx = None

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".markdown", ".csv"}


@dataclass
class RawPage:
    """A single unit of extracted text before chunking (a PDF page, a whole
    txt/md file, or a CSV represented as text)."""
    document_name: str
    page_number: int  # 1-indexed. None/0 for docs without native pages.
    text: str


@dataclass
class LoadedDocument:
    document_name: str
    file_path: str
    file_type: str
    file_hash: str
    pages: List[RawPage] = field(default_factory=list)


def _hash_file(path: str) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()


def _load_pdf(path: str, name: str) -> List[RawPage]:
    pages = []
    if PdfReader is None:
        raise RuntimeError("pypdf is required to read PDF files")
    reader = PdfReader(path)
    for i, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        pages.append(RawPage(document_name=name, page_number=i, text=text))
    return pages


def _load_docx(path: str, name: str) -> List[RawPage]:
    if docx is None:
        raise RuntimeError("python-docx is required to read DOCX files")
    d = docx.Document(path)
    full_text = "\n".join(p.text for p in d.paragraphs)
    # DOCX has no reliable native page boundaries via python-docx, so the
    # whole document is treated as a single logical "page" (page_number=None)
    return [RawPage(document_name=name, page_number=None, text=full_text)]


def _load_txt_or_md(path: str, name: str) -> List[RawPage]:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()
    return [RawPage(document_name=name, page_number=None, text=text)]


def _load_csv(path: str, name: str) -> List[RawPage]:
    rows_text = []
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        reader = csv.reader(f)
        header = next(reader, None)
        for row in reader:
            if header:
                pairs = [f"{h}: {v}" for h, v in zip(header, row)]
                rows_text.append(" | ".join(pairs))
            else:
                rows_text.append(" | ".join(row))
    text = "\n".join(rows_text)
    return [RawPage(document_name=name, page_number=None, text=text)]


def load_document(path: str) -> LoadedDocument:
    """Load a single document of any supported type."""
    name = os.path.basename(path)
    ext = os.path.splitext(path)[1].lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {ext}")

    if ext == ".pdf":
        pages = _load_pdf(path, name)
    elif ext == ".docx":
        pages = _load_docx(path, name)
    elif ext in (".txt", ".md", ".markdown"):
        pages = _load_txt_or_md(path, name)
    elif ext == ".csv":
        pages = _load_csv(path, name)
    else:
        raise ValueError(f"Unsupported file type: {ext}")

    return LoadedDocument(
        document_name=name,
        file_path=path,
        file_type=ext.lstrip("."),
        file_hash=_hash_file(path),
        pages=pages,
    )


def scan_folder(folder: str) -> List[str]:
    """Return all supported file paths inside a folder (recursively)."""
    found = []
    for root, _, files in os.walk(folder):
        for fname in files:
            if os.path.splitext(fname)[1].lower() in SUPPORTED_EXTENSIONS:
                found.append(os.path.join(root, fname))
    return sorted(found)


def load_new_or_changed_documents(folder: str, index_state: Dict[str, str]) -> List[LoadedDocument]:
    """
    Compare files on disk against `index_state` (a dict of file_path -> file_hash
    coming from the previously persisted index). Only new or modified files are
    loaded and returned, which is what allows incremental re-indexing without
    rebuilding the whole vector store.
    """
    to_load = []
    for path in scan_folder(folder):
        current_hash = _hash_file(path)
        if index_state.get(path) != current_hash:
            to_load.append(path)

    docs = []
    for path in to_load:
        try:
            docs.append(load_document(path))
        except Exception as e:
            print(f"[loader] Failed to load {path}: {e}")
    return docs
